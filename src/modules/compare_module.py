"""
Module de comparaison de données
Compare des données Excel avec d'autres fichiers Excel ou PDF/Word
"""

import customtkinter as ctk
from tkinter import filedialog, messagebox
from typing import Dict, Any, Optional, List, Set
import pandas as pd
from pathlib import Path
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading

from .base_module import BaseModule
from ..ui.components.tooltip import Tooltip
from ..ui.components.file_selector import FileSelector
from ..ui.components.preview_table import PreviewTable
from ..ui.components.stat_card import StatCardGroup
from ..core.constants import COLORS
from ..utils.excel_utils import ExcelUtils


class CompareModule(BaseModule):
    """
    Module de comparaison de données entre fichiers

    Fonctionnalités:
    - Comparaison Excel ↔ Excel
    - Comparaison Excel ↔ PDF/Word (extraction de texte)
    - Recherche exacte ou approximative (fuzzy matching)
    - Statistiques détaillées des correspondances
    - Export des résultats avec surbrillance
    """

    MODULE_ID = "compare"
    MODULE_NAME = "Comparaison de données"
    MODULE_DESCRIPTION = "Compare des données entre fichiers Excel, PDF ou Word"
    MODULE_ICON = "⚖️"

    def __init__(self, *args, **kwargs):
        self._stop_event = threading.Event()
        self._executor: Optional[ThreadPoolExecutor] = None
        super().__init__(*args, **kwargs)

    def _create_interface(self):
        """Crée l'interface du module"""
        # Onglets pour les différents modes de comparaison
        self.tabview = ctk.CTkTabview(self.frame, fg_color=COLORS["bg_card"])
        self.tabview.pack(fill="both", expand=True, padx=10, pady=10)

        self.tabview.add("Excel ↔ Excel")
        self.tabview.add("Excel ↔ Document")
        self.tabview.add("Résultats")

        self._create_excel_excel_tab()
        self._create_excel_document_tab()
        self._create_results_tab()

        # Variables pour les résultats
        self.df_found: Optional[pd.DataFrame] = None
        self.df_not_found: Optional[pd.DataFrame] = None

    def _create_excel_excel_tab(self):
        """Onglet comparaison Excel ↔ Excel"""
        tab = self.tabview.tab("Excel ↔ Excel")

        scroll = ctk.CTkScrollableFrame(tab, fg_color="transparent")
        scroll.pack(fill="both", expand=True, padx=10, pady=10)

        # === Fichier 1 ===
        section1 = ctk.CTkFrame(scroll, fg_color=COLORS["bg_card"], corner_radius=10)
        section1.pack(fill="x", pady=(0, 10))

        ctk.CTkLabel(
            section1, text="📄 FICHIER 1 (Source)", font=("Segoe UI", 12, "bold")
        ).pack(anchor="w", padx=15, pady=(15, 10))

        self.file1_selector = FileSelector(
            section1,
            label="Fichier Excel",
            tooltip="Fichier contenant les données à rechercher",
            on_file_loaded=self._on_file1_loaded
        )
        self.file1_selector.pack(fill="x", padx=15, pady=(0, 10))

        col1_frame = ctk.CTkFrame(section1, fg_color="transparent")
        col1_frame.pack(fill="x", padx=15, pady=(0, 15))

        ctk.CTkLabel(col1_frame, text="Colonne de recherche:", width=150).pack(side="left")
        self.col1_combo = ctk.CTkComboBox(
            col1_frame, values=["(Charger fichier)"], state="disabled", width=200
        )
        self.col1_combo.pack(side="left", padx=(10, 0))
        Tooltip(self.col1_combo, "Colonne contenant les valeurs à rechercher")

        # === Fichier 2 ===
        section2 = ctk.CTkFrame(scroll, fg_color=COLORS["bg_card"], corner_radius=10)
        section2.pack(fill="x", pady=(0, 10))

        ctk.CTkLabel(
            section2, text="📋 FICHIER 2 (Référence)", font=("Segoe UI", 12, "bold")
        ).pack(anchor="w", padx=15, pady=(15, 10))

        self.file2_selector = FileSelector(
            section2,
            label="Fichier Excel",
            tooltip="Fichier dans lequel rechercher les correspondances",
            on_file_loaded=self._on_file2_loaded
        )
        self.file2_selector.pack(fill="x", padx=15, pady=(0, 10))

        col2_frame = ctk.CTkFrame(section2, fg_color="transparent")
        col2_frame.pack(fill="x", padx=15, pady=(0, 15))

        ctk.CTkLabel(col2_frame, text="Colonne de référence:", width=150).pack(side="left")
        self.col2_combo = ctk.CTkComboBox(
            col2_frame, values=["(Charger fichier)"], state="disabled", width=200
        )
        self.col2_combo.pack(side="left", padx=(10, 0))

        # === Options ===
        section3 = ctk.CTkFrame(scroll, fg_color=COLORS["bg_card"], corner_radius=10)
        section3.pack(fill="x", pady=(0, 10))

        ctk.CTkLabel(
            section3, text="⚙️ OPTIONS", font=("Segoe UI", 12, "bold")
        ).pack(anchor="w", padx=15, pady=(15, 10))

        options_frame = ctk.CTkFrame(section3, fg_color="transparent")
        options_frame.pack(fill="x", padx=15, pady=(0, 15))

        # Mode de recherche
        self.exact_match_var = ctk.BooleanVar(value=True)
        ctk.CTkRadioButton(
            options_frame, text="Correspondance exacte",
            variable=self.exact_match_var, value=True
        ).pack(side="left", padx=(0, 20))

        ctk.CTkRadioButton(
            options_frame, text="Correspondance approximative",
            variable=self.exact_match_var, value=False
        ).pack(side="left")

        # Options supplémentaires
        options2_frame = ctk.CTkFrame(section3, fg_color="transparent")
        options2_frame.pack(fill="x", padx=15, pady=(0, 15))

        self.case_sensitive_var = ctk.BooleanVar(value=False)
        cb1 = ctk.CTkCheckBox(
            options2_frame, text="Sensible à la casse",
            variable=self.case_sensitive_var
        )
        cb1.pack(side="left", padx=(0, 20))

        # Seuil de similarité
        ctk.CTkLabel(options2_frame, text="Seuil similarité:").pack(side="left", padx=(20, 5))
        self.similarity_slider = ctk.CTkSlider(
            options2_frame, from_=50, to=100, number_of_steps=50, width=150
        )
        self.similarity_slider.set(80)
        self.similarity_slider.pack(side="left")

        self.similarity_label = ctk.CTkLabel(options2_frame, text="80%", width=40)
        self.similarity_label.pack(side="left", padx=5)

        self.similarity_slider.configure(command=self._update_similarity_label)

        # === Boutons ===
        btn_frame = ctk.CTkFrame(scroll, fg_color="transparent")
        btn_frame.pack(fill="x", pady=10)

        self.compare_excel_btn = ctk.CTkButton(
            btn_frame,
            text="⚖️ COMPARER LES FICHIERS",
            font=("Segoe UI", 14, "bold"),
            height=50,
            fg_color=COLORS["accent_primary"],
            command=self._compare_excel_files
        )
        self.compare_excel_btn.pack(fill="x", pady=(0, 10))

        # Progression
        self.progress_bar_excel = ctk.CTkProgressBar(btn_frame, height=8)
        self.progress_bar_excel.pack(fill="x")
        self.progress_bar_excel.set(0)

        self.status_label_excel = ctk.CTkLabel(
            btn_frame, text="Prêt", font=("Segoe UI", 10), text_color=COLORS["text_muted"]
        )
        self.status_label_excel.pack(anchor="w", pady=(5, 0))

    def _create_excel_document_tab(self):
        """Onglet comparaison Excel ↔ PDF/Word"""
        tab = self.tabview.tab("Excel ↔ Document")

        scroll = ctk.CTkScrollableFrame(tab, fg_color="transparent")
        scroll.pack(fill="both", expand=True, padx=10, pady=10)

        # === Fichier Excel ===
        section1 = ctk.CTkFrame(scroll, fg_color=COLORS["bg_card"], corner_radius=10)
        section1.pack(fill="x", pady=(0, 10))

        ctk.CTkLabel(
            section1, text="📊 FICHIER EXCEL", font=("Segoe UI", 12, "bold")
        ).pack(anchor="w", padx=15, pady=(15, 10))

        self.excel_doc_selector = FileSelector(
            section1,
            label="Fichier Excel",
            tooltip="Fichier contenant les données à rechercher",
            on_file_loaded=self._on_excel_doc_loaded
        )
        self.excel_doc_selector.pack(fill="x", padx=15, pady=(0, 10))

        col_frame = ctk.CTkFrame(section1, fg_color="transparent")
        col_frame.pack(fill="x", padx=15, pady=(0, 15))

        ctk.CTkLabel(col_frame, text="Colonne de recherche:", width=150).pack(side="left")
        self.col_doc_combo = ctk.CTkComboBox(
            col_frame, values=["(Charger fichier)"], state="disabled", width=200
        )
        self.col_doc_combo.pack(side="left", padx=(10, 0))

        # === Fichier Document ===
        section2 = ctk.CTkFrame(scroll, fg_color=COLORS["bg_card"], corner_radius=10)
        section2.pack(fill="x", pady=(0, 10))

        ctk.CTkLabel(
            section2, text="📄 DOCUMENT (PDF/Word)", font=("Segoe UI", 12, "bold")
        ).pack(anchor="w", padx=15, pady=(15, 10))

        doc_frame = ctk.CTkFrame(section2, fg_color="transparent")
        doc_frame.pack(fill="x", padx=15, pady=(0, 15))

        self.doc_path_var = ctk.StringVar()
        ctk.CTkEntry(
            doc_frame, textvariable=self.doc_path_var, width=300, state="disabled"
        ).pack(side="left", padx=(0, 10))

        ctk.CTkButton(
            doc_frame, text="📁 Parcourir", width=100,
            command=self._browse_document
        ).pack(side="left")

        self.doc_info_label = ctk.CTkLabel(
            section2, text="", text_color=COLORS["text_muted"]
        )
        self.doc_info_label.pack(anchor="w", padx=15, pady=(0, 15))

        # === Boutons ===
        btn_frame = ctk.CTkFrame(scroll, fg_color="transparent")
        btn_frame.pack(fill="x", pady=10)

        self.compare_doc_btn = ctk.CTkButton(
            btn_frame,
            text="⚖️ COMPARER EXCEL ↔ DOCUMENT",
            font=("Segoe UI", 14, "bold"),
            height=50,
            fg_color=COLORS["accent_primary"],
            command=self._compare_excel_document
        )
        self.compare_doc_btn.pack(fill="x", pady=(0, 10))

        # Progression
        self.progress_bar_doc = ctk.CTkProgressBar(btn_frame, height=8)
        self.progress_bar_doc.pack(fill="x")
        self.progress_bar_doc.set(0)

        self.status_label_doc = ctk.CTkLabel(
            btn_frame, text="Prêt", font=("Segoe UI", 10), text_color=COLORS["text_muted"]
        )
        self.status_label_doc.pack(anchor="w", pady=(5, 0))

    def _create_results_tab(self):
        """Onglet des résultats"""
        tab = self.tabview.tab("Résultats")

        # Statistiques
        self.stats_cards = StatCardGroup(tab)
        self.stats_cards.pack(fill="x", padx=10, pady=10)

        self.stats_cards.add_card("total", "Total recherché", "0", "📊")
        self.stats_cards.add_card("found", "Trouvées", "0", "✓", color=COLORS["success"])
        self.stats_cards.add_card("not_found", "Non trouvées", "0", "✗", color=COLORS["error"])
        self.stats_cards.add_card("rate", "Taux", "0%", "📈", color=COLORS["info"])

        # Sous-onglets pour les résultats
        self.results_tabview = ctk.CTkTabview(tab, fg_color="transparent")
        self.results_tabview.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        self.results_tabview.add("✓ Trouvées")
        self.results_tabview.add("✗ Non trouvées")

        self.found_table = PreviewTable(
            self.results_tabview.tab("✓ Trouvées"),
            title="Correspondances trouvées",
            max_rows=1000
        )
        self.found_table.pack(fill="both", expand=True, padx=5, pady=5)

        self.not_found_table = PreviewTable(
            self.results_tabview.tab("✗ Non trouvées"),
            title="Valeurs non trouvées",
            max_rows=1000
        )
        self.not_found_table.pack(fill="both", expand=True, padx=5, pady=5)

        # Boutons d'export
        export_frame = ctk.CTkFrame(tab, fg_color="transparent")
        export_frame.pack(fill="x", padx=10, pady=(0, 10))

        self.export_found_btn = ctk.CTkButton(
            export_frame, text="📤 Exporter trouvées", width=150,
            fg_color=COLORS["success"], state="disabled",
            command=lambda: self._export_results("found")
        )
        self.export_found_btn.pack(side="left", padx=(0, 10))

        self.export_not_found_btn = ctk.CTkButton(
            export_frame, text="📤 Exporter non trouvées", width=150,
            fg_color=COLORS["error"], state="disabled",
            command=lambda: self._export_results("not_found")
        )
        self.export_not_found_btn.pack(side="left", padx=(0, 10))

        self.export_all_btn = ctk.CTkButton(
            export_frame, text="📤 Exporter tout", width=150,
            fg_color=COLORS["accent_primary"], state="disabled",
            command=lambda: self._export_results("all")
        )
        self.export_all_btn.pack(side="left")

    def _update_similarity_label(self, value):
        """Met à jour le label du seuil de similarité"""
        self.similarity_label.configure(text=f"{int(value)}%")

    def _on_file1_loaded(self, df: pd.DataFrame):
        """Callback fichier 1 chargé"""
        columns = list(df.columns)
        self.col1_combo.configure(state="normal", values=columns)
        if columns:
            self.col1_combo.set(columns[0])
        self.log_info(f"Fichier 1 chargé: {len(df)} lignes")

    def _on_file2_loaded(self, df: pd.DataFrame):
        """Callback fichier 2 chargé"""
        columns = list(df.columns)
        self.col2_combo.configure(state="normal", values=columns)
        if columns:
            self.col2_combo.set(columns[0])
        self.log_info(f"Fichier 2 chargé: {len(df)} lignes")

    def _on_excel_doc_loaded(self, df: pd.DataFrame):
        """Callback fichier Excel (mode document) chargé"""
        columns = list(df.columns)
        self.col_doc_combo.configure(state="normal", values=columns)
        if columns:
            self.col_doc_combo.set(columns[0])
        self.log_info(f"Fichier Excel chargé: {len(df)} lignes")

    def _browse_document(self):
        """Sélectionne un document PDF ou Word"""
        filepath = filedialog.askopenfilename(
            title="Sélectionner un document",
            filetypes=[
                ("Documents", "*.pdf *.docx *.doc"),
                ("PDF", "*.pdf"),
                ("Word", "*.docx *.doc")
            ]
        )
        if filepath:
            self.doc_path_var.set(filepath)
            ext = Path(filepath).suffix.lower()
            self.doc_info_label.configure(
                text=f"Type: {ext.upper()[1:]} - {Path(filepath).name}",
                text_color=COLORS["success"]
            )

    def _extract_text_from_document(self, filepath: str) -> str:
        """Extrait le texte d'un document PDF ou Word"""
        ext = Path(filepath).suffix.lower()
        text = ""

        if ext == ".pdf":
            try:
                import PyPDF2
                with open(filepath, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    for page in reader.pages:
                        text += page.extract_text() + "\n"
            except ImportError:
                # Fallback avec pdfplumber si disponible
                try:
                    import pdfplumber
                    with pdfplumber.open(filepath) as pdf:
                        for page in pdf.pages:
                            text += (page.extract_text() or "") + "\n"
                except ImportError:
                    raise ImportError("PyPDF2 ou pdfplumber requis pour lire les PDF")

        elif ext in [".docx", ".doc"]:
            try:
                from docx import Document
                doc = Document(filepath)
                for para in doc.paragraphs:
                    text += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + " "
                    text += "\n"
            except ImportError:
                raise ImportError("python-docx requis pour lire les fichiers Word")

        return text

    def _compare_excel_files(self):
        """Compare deux fichiers Excel"""
        if not self.file1_selector.is_loaded() or not self.file2_selector.is_loaded():
            messagebox.showwarning("Attention", "Veuillez charger les deux fichiers")
            return

        self._stop_event.clear()
        self.compare_excel_btn.configure(state="disabled")
        self.progress_bar_excel.set(0)
        self.status_label_excel.configure(text="Comparaison en cours...")

        # Lancer dans un thread
        thread = threading.Thread(target=self._do_excel_comparison, daemon=True)
        thread.start()

    def _do_excel_comparison(self):
        """Effectue la comparaison Excel (dans un thread)"""
        try:
            df1 = self.file1_selector.get_dataframe()
            df2 = self.file2_selector.get_dataframe()
            col1 = self.col1_combo.get()
            col2 = self.col2_combo.get()
            exact = self.exact_match_var.get()
            case_sensitive = self.case_sensitive_var.get()
            threshold = self.similarity_slider.get() / 100

            # Préparer les valeurs
            values_to_find = df1[col1].dropna().astype(str).unique()
            reference_values = set(df2[col2].dropna().astype(str))

            if not case_sensitive:
                reference_values = {v.lower() for v in reference_values}

            found_values = []
            not_found_values = []
            total = len(values_to_find)

            for idx, value in enumerate(values_to_find):
                if self._stop_event.is_set():
                    break

                search_value = value if case_sensitive else value.lower()

                if exact:
                    is_found = search_value in reference_values
                else:
                    # Correspondance approximative avec difflib
                    from difflib import SequenceMatcher
                    is_found = any(
                        SequenceMatcher(None, search_value, ref).ratio() >= threshold
                        for ref in reference_values
                    )

                if is_found:
                    found_values.append(value)
                else:
                    not_found_values.append(value)

                # Mise à jour progression
                progress = (idx + 1) / total
                self.frame.after(0, lambda p=progress: self.progress_bar_excel.set(p))

            # Créer les DataFrames de résultats
            self.df_found = df1[df1[col1].astype(str).isin(found_values)]
            self.df_not_found = df1[df1[col1].astype(str).isin(not_found_values)]

            # Mettre à jour l'interface
            self.frame.after(0, lambda: self._display_comparison_results(total))

        except Exception as e:
            self.frame.after(0, lambda: messagebox.showerror("Erreur", str(e)))
            self.log_error(str(e))

        finally:
            self.frame.after(0, lambda: self.compare_excel_btn.configure(state="normal"))

    def _compare_excel_document(self):
        """Compare Excel avec un document PDF/Word"""
        if not self.excel_doc_selector.is_loaded():
            messagebox.showwarning("Attention", "Veuillez charger un fichier Excel")
            return

        doc_path = self.doc_path_var.get()
        if not doc_path:
            messagebox.showwarning("Attention", "Veuillez sélectionner un document")
            return

        self._stop_event.clear()
        self.compare_doc_btn.configure(state="disabled")
        self.progress_bar_doc.set(0)
        self.status_label_doc.configure(text="Extraction du texte...")

        thread = threading.Thread(target=self._do_document_comparison, daemon=True)
        thread.start()

    def _do_document_comparison(self):
        """Effectue la comparaison avec document (dans un thread)"""
        try:
            doc_path = self.doc_path_var.get()

            # Extraire le texte du document
            doc_text = self._extract_text_from_document(doc_path)
            doc_text_lower = doc_text.lower()

            self.frame.after(0, lambda: self.status_label_doc.configure(text="Comparaison en cours..."))

            df = self.excel_doc_selector.get_dataframe()
            col = self.col_doc_combo.get()

            values_to_find = df[col].dropna().astype(str).unique()
            found_values = []
            not_found_values = []
            total = len(values_to_find)

            for idx, value in enumerate(values_to_find):
                if self._stop_event.is_set():
                    break

                # Recherche insensible à la casse par défaut
                if value.lower() in doc_text_lower:
                    found_values.append(value)
                else:
                    not_found_values.append(value)

                progress = (idx + 1) / total
                self.frame.after(0, lambda p=progress: self.progress_bar_doc.set(p))

            self.df_found = df[df[col].astype(str).isin(found_values)]
            self.df_not_found = df[df[col].astype(str).isin(not_found_values)]

            self.frame.after(0, lambda: self._display_comparison_results(total))

        except Exception as e:
            self.frame.after(0, lambda: messagebox.showerror("Erreur", str(e)))
            self.log_error(str(e))

        finally:
            self.frame.after(0, lambda: self.compare_doc_btn.configure(state="normal"))

    def _display_comparison_results(self, total: int):
        """Affiche les résultats de comparaison"""
        found = len(self.df_found) if self.df_found is not None else 0
        not_found = len(self.df_not_found) if self.df_not_found is not None else 0
        rate = (found / total * 100) if total > 0 else 0

        # Mettre à jour les statistiques
        self.stats_cards.update_card("total", str(total))
        self.stats_cards.update_card("found", str(found))
        self.stats_cards.update_card("not_found", str(not_found))
        self.stats_cards.update_card("rate", f"{rate:.1f}%")

        # Charger les tables
        if self.df_found is not None:
            self.found_table.load_data(self.df_found)
        if self.df_not_found is not None:
            self.not_found_table.load_data(self.df_not_found)

        # Activer les boutons d'export
        self.export_found_btn.configure(state="normal" if found > 0 else "disabled")
        self.export_not_found_btn.configure(state="normal" if not_found > 0 else "disabled")
        self.export_all_btn.configure(state="normal" if total > 0 else "disabled")

        # Aller à l'onglet résultats
        self.tabview.set("Résultats")

        # Mise à jour des statuts
        self.status_label_excel.configure(text=f"Terminé: {found}/{total} trouvées", text_color=COLORS["success"])
        self.status_label_doc.configure(text=f"Terminé: {found}/{total} trouvées", text_color=COLORS["success"])

        self.log_success(f"Comparaison terminée: {found}/{total} correspondances ({rate:.1f}%)")

    def _export_results(self, mode: str):
        """Exporte les résultats"""
        filepath = filedialog.asksaveasfilename(
            title="Enregistrer les résultats",
            defaultextension=".xlsx",
            filetypes=[("Fichiers Excel", "*.xlsx")]
        )

        if not filepath:
            return

        try:
            from openpyxl import Workbook
            from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

            wb = Workbook()
            wb.remove(wb.active)

            # Styles
            header_fill = PatternFill(start_color="2E5090", end_color="2E5090", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF")
            found_fill = PatternFill(start_color="C6EFCE", end_color="C6EFCE", fill_type="solid")
            not_found_fill = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")
            border = Border(
                left=Side(style='thin'), right=Side(style='thin'),
                top=Side(style='thin'), bottom=Side(style='thin')
            )

            if mode in ["found", "all"] and self.df_found is not None and len(self.df_found) > 0:
                ws = wb.create_sheet("Trouvées")
                self._write_df_to_sheet(ws, self.df_found, header_fill, header_font, found_fill, border)

            if mode in ["not_found", "all"] and self.df_not_found is not None and len(self.df_not_found) > 0:
                ws = wb.create_sheet("Non_Trouvées")
                self._write_df_to_sheet(ws, self.df_not_found, header_fill, header_font, not_found_fill, border)

            wb.save(filepath)
            messagebox.showinfo("Succès", f"Résultats exportés vers:\n{filepath}")
            self.log_success(f"Export réussi: {filepath}")

        except Exception as e:
            messagebox.showerror("Erreur", str(e))
            self.log_error(str(e))

    def _write_df_to_sheet(self, ws, df, header_fill, header_font, row_fill, border):
        """Écrit un DataFrame dans une feuille avec formatage"""
        # En-têtes
        for col_idx, col_name in enumerate(df.columns, 1):
            cell = ws.cell(row=1, column=col_idx, value=col_name)
            cell.fill = header_fill
            cell.font = header_font
            cell.border = border

        # Données
        for row_idx, row in enumerate(df.itertuples(index=False), 2):
            for col_idx, value in enumerate(row, 1):
                cell = ws.cell(row=row_idx, column=col_idx, value=value)
                cell.fill = row_fill
                cell.border = border

        # Ajuster les colonnes
        for col_idx, col_name in enumerate(df.columns, 1):
            max_length = max(len(str(col_name)), df[col_name].astype(str).str.len().max())
            ws.column_dimensions[ws.cell(row=1, column=col_idx).column_letter].width = min(max_length + 2, 50)

    def validate_inputs(self) -> tuple[bool, str]:
        return True, ""

    def _execute_task(self) -> Dict[str, Any]:
        return {}

    def update_status(self, message: str, level: str = "info"):
        pass

    def update_progress(self, progress: float):
        pass

    def reset(self):
        super().reset()
        self.df_found = None
        self.df_not_found = None
        self.file1_selector.reset()
        self.file2_selector.reset()
        self.excel_doc_selector.reset()
        self.found_table.clear()
        self.not_found_table.clear()
        self.stats_cards.reset_all()
